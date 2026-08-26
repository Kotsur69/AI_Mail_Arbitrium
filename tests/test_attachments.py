"""Reading the files suppliers attach, and reporting the ones we cannot read.

Every fixture is built in memory from real libraries writing real formats --
python-docx writes the .docx, openpyxl the .xlsx, and the PDFs are assembled
byte by byte. No checked-in binaries, so nothing here can drift out of sync
with the parsers it is meant to exercise.

The failure cases carry as much weight as the successes. A scan and an empty
envelope both produce no text, and the whole point of the error field is that
the person working the review queue can tell those two apart.
"""

from io import BytesIO

import pytest

from arbitrium.attachments.base import MAX_ATTACHMENT_BYTES, Attachment, Extraction, Kind
from arbitrium.attachments.extract import (
    TRUNCATED,
    attachment_text,
    extract,
    extract_all,
    truncate,
)


def pdf_with_text(text: str) -> bytes:
    """A minimal but structurally valid PDF carrying one line of text."""
    content = f"BT /F1 24 Tf 72 700 Td ({text}) Tj ET".encode("latin-1")
    objects = [
        b"<</Type/Catalog/Pages 2 0 R>>",
        b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]"
        b"/Resources<</Font<</F1 4 0 R>>>>/Contents 5 0 R>>",
        b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
        b"<</Length %d>>stream\n%s\nendstream" % (len(content), content),
    ]

    out = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for number, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += b"%d 0 obj" % number + body + b"endobj\n"

    start_xref = len(out)
    out += b"xref\n0 %d\n0000000000 65535 f \n" % (len(objects) + 1)
    for offset in offsets:
        out += b"%010d 00000 n \n" % offset
    out += b"trailer<</Size %d/Root 1 0 R>>\nstartxref\n%d\n%%%%EOF\n" % (
        len(objects) + 1,
        start_xref,
    )
    return bytes(out)


def docx_bytes(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    import docx

    document = docx.Document()
    for text in paragraphs:
        document.add_paragraph(text)

    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row, values in zip(table.rows, table_rows):
            for cell, value in zip(row.cells, values):
                cell.text = value

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def xlsx_bytes(sheet_name: str, rows: list[list[object]]) -> bytes:
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = sheet_name
    for row in rows:
        sheet.append(row)

    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


# --- naming: what a file claims to be -------------------------------------


@pytest.mark.parametrize(
    ("filename", "extension"),
    [
        ("zgoda.pdf", "pdf"),
        ("ZGODA.PDF", "pdf"),
        ("  aneks.DocX  ", "docx"),
        ("archiwum.tar.gz", "gz"),
        ("README", ""),
        ("kropka.", ""),
        ("", ""),
    ],
)
def test_extension_is_the_last_suffix_lowercased(filename: str, extension: str) -> None:
    assert Attachment(filename).extension == extension


@pytest.mark.parametrize(
    ("filename", "kind"),
    [
        ("zgoda.pdf", Kind.PDF),
        ("aneks.docx", Kind.DOCX),
        ("cennik.xlsx", Kind.XLSX),
        ("notatka.txt", Kind.TEXT),
        ("skan.jpg", Kind.IMAGE),
        ("umowa.doc", Kind.LEGACY_OFFICE),
        ("paczka.zip", Kind.UNSUPPORTED),
    ],
)
def test_kind_follows_the_extension(filename: str, kind: Kind) -> None:
    assert Attachment(filename).kind is kind


def test_only_readable_kinds_are_worth_loading() -> None:
    # A signature logo travels on half the mail in a thread; its bytes buy nothing.
    assert Attachment("logo.png").is_extractable is False
    assert Attachment("umowa.doc").is_extractable is False
    assert Attachment("zgoda.pdf").is_extractable is True


def test_an_oversized_file_is_not_loaded_however_readable_its_type() -> None:
    assert Attachment("katalog.pdf", size_bytes=MAX_ATTACHMENT_BYTES + 1).is_extractable is False
    assert Attachment("katalog.pdf", size_bytes=MAX_ATTACHMENT_BYTES).is_extractable is True


def test_unknown_size_is_not_treated_as_too_big() -> None:
    # Outlook does not always answer for Size, and "we don't know" must not
    # silently mean "skip it" -- that would drop consent without saying so.
    assert Attachment("zgoda.pdf", size_bytes=None).is_extractable is True


# --- extraction: what a file actually says ---------------------------------


def test_pdf_text_is_extracted() -> None:
    result = extract(Attachment("zgoda.pdf", pdf_with_text("Wyrazamy zgode na przetwarzanie")))

    assert result.error is None
    assert result.kind is Kind.PDF
    assert "Wyrazamy zgode" in result.text


def test_pdf_without_a_text_layer_points_at_a_vision_model() -> None:
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buffer = BytesIO()
    writer.write(buffer)

    result = extract(Attachment("skan.pdf", buffer.getvalue()))

    # This is the scanned-signature case: not a broken file, a different tool.
    assert not result.has_text
    assert "wizyjnego" in (result.error or "")


def test_encrypted_pdf_is_reported_as_a_password_not_as_damage() -> None:
    from pypdf import PdfWriter

    writer = PdfWriter(clone_from=BytesIO(pdf_with_text("tajne")))
    writer.encrypt("haslo")
    buffer = BytesIO()
    writer.write(buffer)

    result = extract(Attachment("zgoda.pdf", buffer.getvalue()))

    assert not result.has_text
    assert "zaszyfrowany" in (result.error or "")


def test_docx_paragraphs_are_read() -> None:
    data = docx_bytes(["Szanowni Panstwo,", "Akceptujemy warunki."])

    result = extract(Attachment("odpowiedz.docx", data))

    assert "Akceptujemy warunki." in result.text


def test_docx_tables_are_read_because_consent_is_often_a_grid() -> None:
    data = docx_bytes(["Formularz"], table_rows=[["Wyrazam zgode", "TAK"]])

    result = extract(Attachment("formularz.docx", data))

    assert "Wyrazam zgode | TAK" in result.text


def test_xlsx_rows_are_read_and_labelled_by_sheet() -> None:
    data = xlsx_bytes("Zgody", [["Dostawca", "Status"], ["Firma sp. z o.o.", "zgoda"]])

    result = extract(Attachment("zestawienie.xlsx", data))

    assert "[Zgody]" in result.text
    assert "Firma sp. z o.o. | zgoda" in result.text


def test_empty_spreadsheet_rows_are_dropped() -> None:
    data = xlsx_bytes("Arkusz", [["Dostawca"], [None, None], ["Firma"]])

    lines = extract(Attachment("puste.xlsx", data)).text.splitlines()

    assert lines == ["[Arkusz]", "Dostawca", "Firma"]


def test_windows_encoded_text_survives() -> None:
    # Polish .txt attachments are written on Windows far more often than in UTF-8.
    data = "Wyrażamy zgodę.".encode("cp1250")

    result = extract(Attachment("odpowiedz.txt", data))

    assert result.text == "Wyrażamy zgodę."


def test_utf8_text_with_a_byte_order_mark_loses_the_mark() -> None:
    result = extract(Attachment("notatka.txt", "Zgoda.".encode("utf-8-sig")))

    assert result.text == "Zgoda."


# --- failure: what a file could not say ------------------------------------


def test_an_image_is_reported_not_skipped() -> None:
    result = extract(Attachment("podpis.png", b"\x89PNG\r\n\x1a\n"))

    assert result == Extraction("podpis.png", Kind.IMAGE, error="obraz -- wymaga modelu wizyjnego")


def test_a_legacy_binary_format_names_itself_as_the_reason() -> None:
    result = extract(Attachment("umowa.doc", b"\xd0\xcf\x11\xe0"))

    assert "stary format" in (result.error or "")


def test_an_unloaded_attachment_says_the_content_is_missing() -> None:
    # The name came back from Outlook but the bytes did not: distinguishable
    # from a file we read and found empty.
    result = extract(Attachment("zgoda.pdf", data=None))

    assert result.error == "tresc nie zostala wczytana"


def test_a_corrupt_file_returns_a_reason_rather_than_raising() -> None:
    result = extract(Attachment("zgoda.pdf", b"to nie jest PDF"))

    assert not result.has_text
    assert result.error


def test_a_file_holding_only_whitespace_counts_as_empty() -> None:
    result = extract(Attachment("puste.txt", b"   \n\n   "))

    assert result.error == "plik nie zawiera tekstu"


# --- assembling the prompt --------------------------------------------------


def test_truncation_says_that_it_happened() -> None:
    cut = truncate("a" * 100, limit=10)

    assert cut.startswith("aaaa")
    assert cut.endswith(TRUNCATED)


def test_text_within_the_limit_is_untouched() -> None:
    assert truncate("krotki tekst", limit=100) == "krotki tekst"


def test_each_readable_file_is_labelled_by_name() -> None:
    extractions = [
        Extraction("zgoda.pdf", Kind.PDF, text="Wyrazamy zgode."),
        Extraction("cennik.xlsx", Kind.XLSX, text="[Arkusz]\n100 | 200"),
    ]

    block = attachment_text(extractions)

    assert "--- ZALACZNIK: zgoda.pdf ---\nWyrazamy zgode." in block
    assert "--- ZALACZNIK: cennik.xlsx ---" in block


def test_unreadable_files_contribute_nothing_to_the_prompt() -> None:
    extractions = [
        Extraction("skan.pdf", Kind.PDF, error="skan"),
        Extraction("zgoda.pdf", Kind.PDF, text="Zgoda."),
    ]

    block = attachment_text(extractions)

    assert "skan.pdf" not in block
    assert "Zgoda." in block


def test_no_readable_files_means_no_prompt_at_all() -> None:
    assert attachment_text([Extraction("skan.pdf", Kind.PDF, error="skan")]) == ""
    assert attachment_text([]) == ""


def test_extract_all_keeps_the_order_it_was_given() -> None:
    results = extract_all(
        (
            Attachment("a.txt", b"pierwszy"),
            Attachment("b.png", b""),
            Attachment("c.txt", b"trzeci"),
        )
    )

    assert [r.filename for r in results] == ["a.txt", "b.png", "c.txt"]
    assert [r.has_text for r in results] == [True, False, True]
