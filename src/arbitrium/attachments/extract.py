"""Pull plain text out of the attachment formats suppliers actually send.

Everything runs in-process on bytes already in memory: no temp files, no
subprocess, no service. That is a privacy property, not a performance one --
the contract never touches disk outside Outlook's own store.

A format we cannot read is reported, never skipped. Two cases matter enough to
name individually rather than lump into "failed":

  * an encrypted PDF, which is a password nobody has, not a broken file;
  * a PDF with no text layer, which is a scan -- the point where a vision model
    is the only remaining option, and the reason this returns a reason.

That second case now has somewhere to go. Hand `extract` a `VisionReader` and a
scan is transcribed instead of refused; leave it out and the behaviour is
exactly what it was, down to the wording of the error. The transcription is
marked `via_vision` the whole way, because text a model read off pixels is not
the same evidence as text a parser lifted out of a file.
"""

from __future__ import annotations

from io import BytesIO
from typing import TYPE_CHECKING

from arbitrium.attachments.base import Attachment, Extraction, Kind
from arbitrium.normalize import clean_whitespace

if TYPE_CHECKING:
    from arbitrium.attachments.vision import VisionReader

# One signed consent runs a page or two; twenty thousand characters is already
# far past that. Beyond it the text is boilerplate the model pays for by token.
MAX_CHARS_PER_ATTACHMENT = 20_000

# A spreadsheet this long is a price list, not a statement of position.
MAX_XLSX_ROWS = 2_000

TRUNCATED = "\n[... tekst zalacznika obciety ...]"

# Polish attachments arrive in all of these. Order matters: utf-8 first because
# it fails loudly on the others' bytes, cp1250 before iso-8859-2 because Windows
# is what these are written on.
TEXT_ENCODINGS = ("utf-8-sig", "utf-8", "cp1250", "iso-8859-2", "latin-1")


def _pdf(data: bytes) -> str:
    from pypdf import PdfReader  # noqa: PLC0415 - parser imported where it is used

    reader = PdfReader(BytesIO(data))
    if reader.is_encrypted:
        raise ValueError("PDF zaszyfrowany haslem")

    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    if not text.strip():
        raise ValueError("PDF bez warstwy tekstowej (skan?) -- wymaga modelu wizyjnego")
    return text


def _docx(data: bytes) -> str:
    import docx  # noqa: PLC0415

    document = docx.Document(BytesIO(data))
    parts = [p.text for p in document.paragraphs]

    # Consent wording lives in tables at least as often as in prose -- a
    # two-column "Wyrazam zgode / TAK" grid is a normal way to answer.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _xlsx(data: bytes) -> str:
    from openpyxl import load_workbook  # noqa: PLC0415

    workbook = load_workbook(BytesIO(data), read_only=True, data_only=True)
    try:
        lines: list[str] = []
        for sheet in workbook.worksheets:
            lines.append(f"[{sheet.title}]")
            for index, row in enumerate(sheet.iter_rows(values_only=True)):
                if index >= MAX_XLSX_ROWS:
                    lines.append("[... dalsze wiersze pominiete ...]")
                    break
                cells = ["" if value is None else str(value) for value in row]
                # One populated cell anywhere in the sheet widens every row, so
                # trailing blanks are padding, not columns. Left in, they hang a
                # separator off the end of every line the model reads.
                while cells and not cells[-1].strip():
                    cells.pop()
                if cells:
                    lines.append(" | ".join(cells))
        return "\n".join(lines)
    finally:
        workbook.close()


def _text(data: bytes) -> str:
    for encoding in TEXT_ENCODINGS:
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    # latin-1 maps every byte, so reaching here means the tuple lost its anchor.
    raise ValueError("nie udalo sie zdekodowac tekstu")


READERS = {
    Kind.PDF: _pdf,
    Kind.DOCX: _docx,
    Kind.XLSX: _xlsx,
    Kind.TEXT: _text,
}

# Why a kind was never opened, in the words the reviewer sees.
NOT_READ = {
    Kind.IMAGE: "obraz -- wymaga modelu wizyjnego",
    Kind.LEGACY_OFFICE: "stary format binarny (doc/xls/rtf) -- nieobslugiwany",
    Kind.UNSUPPORTED: "nieobslugiwany typ pliku",
}


def truncate(text: str, limit: int = MAX_CHARS_PER_ATTACHMENT) -> str:
    """Cut over-long text, saying so, so a missing quote has an explanation."""
    return text if len(text) <= limit else text[:limit].rstrip() + TRUNCATED


def _transcribe(
    attachment: Attachment, kind: Kind, images: list[bytes], vision: VisionReader
) -> Extraction:
    """Read page images with the vision model, reporting rather than raising.

    A vision model is the most failure-prone thing in this pipeline -- it is the
    one part that needs a second set of weights loaded -- so every way it can go
    wrong ends as a reason on the record. "No VL model loaded in LM Studio" is
    the everyday case, and it must not take the run down with it.
    """
    try:
        raw = vision.transcribe(images)
    except Exception as exc:  # noqa: BLE001 - an absent model is a reason, not a crash
        detail = str(exc) or exc.__class__.__name__
        return Extraction(attachment.filename, kind, error=f"model wizyjny: {detail}")

    text = truncate(clean_whitespace(raw))
    if not text.strip():
        return Extraction(attachment.filename, kind, error="model wizyjny nie odczytal tekstu")
    return Extraction(attachment.filename, kind, text=text, via_vision=True)


def extract(attachment: Attachment, vision: VisionReader | None = None) -> Extraction:
    """Read one attachment. Always returns a record -- text or a stated reason."""
    kind = attachment.kind

    # A scan that arrived as a picture rather than inside a PDF. Checked before
    # the reader table, because there is no parser for it to fall out of, and
    # gated on looking like a page so signature logos never reach the model.
    if kind is Kind.IMAGE and vision is not None and attachment.data is not None:
        from arbitrium.attachments.vision import looks_like_scan  # noqa: PLC0415

        if looks_like_scan(attachment.data):
            return _transcribe(attachment, kind, [attachment.data], vision)

    if kind not in READERS:
        return Extraction(attachment.filename, kind, error=NOT_READ.get(kind, "pominiety"))
    if attachment.data is None:
        return Extraction(attachment.filename, kind, error="tresc nie zostala wczytana")

    try:
        raw = READERS[kind](attachment.data)
    except Exception as exc:  # noqa: BLE001 - a malformed file must not end the run
        # The no-text-layer case lands here, which is exactly where the pixels
        # are still available. An encrypted PDF also lands here and yields no
        # images, so it keeps its own, more useful, error.
        if kind is Kind.PDF and vision is not None:
            from arbitrium.attachments.vision import page_images  # noqa: PLC0415

            images = page_images(attachment.data)
            if images:
                return _transcribe(attachment, kind, images, vision)
        return Extraction(attachment.filename, kind, error=str(exc) or exc.__class__.__name__)

    text = truncate(clean_whitespace(raw))
    if not text.strip():
        return Extraction(attachment.filename, kind, error="plik nie zawiera tekstu")
    return Extraction(attachment.filename, kind, text=text)


def extract_all(
    attachments: tuple[Attachment, ...], vision: VisionReader | None = None
) -> list[Extraction]:
    return [extract(attachment, vision) for attachment in attachments]


def attachment_text(extractions: list[Extraction], limit: int = MAX_CHARS_PER_ATTACHMENT) -> str:
    """The readable attachments as one block, each labelled by filename.

    Labelled because the model is about to be asked what the supplier's position
    is, and "it says yes in the covering note but the annex is a price list" is
    a distinction it can only draw if it can tell the two apart. The label is
    also what a reviewer looks for when checking a quote against its source.
    """
    blocks = [
        f"--- ZALACZNIK: {item.filename} ---\n{item.text}"
        for item in extractions
        if item.has_text
    ]
    return truncate("\n\n".join(blocks), limit)
